import streamlit as st
import pdfplumber
import anthropic
import json
import re
import io

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── CONFIG ────────────────────────────────────────────────────────────────────

st.set_page_config(
    page_title="Deal 1-Pager Generator",
    page_icon="🏢",
    layout="centered"
)

st.markdown("""
<style>
    .block-container { max-width: 680px; padding-top: 2rem; }
</style>
""", unsafe_allow_html=True)

st.title("Deal 1-Pager Generator")
st.caption("Upload a multifamily offering memorandum. Get a standardized 1-page deal summary as an editable Word doc.")

# ── EXTRACTION PROMPT ─────────────────────────────────────────────────────────

EXTRACTION_PROMPT = """You are a commercial real estate analyst. Extract structured data from this multifamily offering memorandum.

Return ONLY valid JSON matching the schema below. If a field is not explicitly stated in the OM, set it to null. Never infer, calculate, or fabricate values not directly written in the OM. For bullet arrays, return concise strings (1-2 sentences each).

Schema:
{
  "deal_name": string or null,
  "address": string or null,
  "city_state": string or null,
  "submarket": string or null,
  "asset_class": string or null,
  "deal_type": string or null,
  "deal_status": string or null,
  "broker": string or null,
  "units": string or null,
  "avg_sf": string or null,
  "year_built": string or null,
  "year_renovated": string or null,
  "physical_occupancy": string or null,
  "purchase_price": string or null,
  "price_per_unit": string or null,
  "going_in_cap_rate": string or null,
  "cap_rate_basis": string or null,
  "investment_thesis": [string],
  "business_plan": [string],
  "construction_type": string or null,
  "parking": string or null,
  "stories": string or null,
  "economic_occupancy": string or null,
  "amenities": string or null,
  "unit_mix": string or null,
  "location_bullets": [string],
  "in_place_rent": string or null,
  "pro_forma_rent": string or null,
  "loss_to_lease": string or null,
  "t12_basis": string or null,
  "t12_egi": string or null,
  "t12_opex": string or null,
  "t12_opex_pct": string or null,
  "t12_noi": string or null,
  "t12_noi_margin": string or null,
  "stab_label": string or null,
  "stab_egi": string or null,
  "stab_opex": string or null,
  "stab_opex_pct": string or null,
  "stab_noi": string or null,
  "stab_noi_margin": string or null,
  "capex_total": string or null,
  "capex_per_unit": string or null,
  "capex_bullets": [string],
  "lender": string or null,
  "debt_type": string or null,
  "term_io": string or null,
  "rate": string or null,
  "ltc_ltv": string or null,
  "equity": string or null,
  "levered_irr": string or null,
  "equity_multiple": string or null,
  "avg_coc": string or null,
  "exit_year": string or null,
  "exit_cap": string or null,
  "tax_notes": string or null,
  "key_risks": [string],
  "why_this_works": [string],
  "broker_process": string or null,
  "guidance": string or null,
  "bid_date": string or null,
  "tour_status": string or null,
  "notes": string or null
}

OM TEXT:
"""

# ── HELPERS ───────────────────────────────────────────────────────────────────

def ns(val, fallback="Not stated"):
    if val is None or val == "" or (isinstance(val, list) and len(val) == 0):
        return fallback
    return str(val)

def add_cell_shading(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)

def remove_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "bottom", "left", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "none")
        tag.set(qn("w:sz"), "0")
        tag.set(qn("w:color"), "FFFFFF")
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def add_para_border_bottom(para, color="CCCCCC", size="4"):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def section_header(cell, text):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    add_para_border_bottom(p, "CCCCCC", "4")
    run = p.add_run(text.upper())
    run.font.name = "Arial"
    run.font.size = Pt(7)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

def bullet(cell, text, bold_label=None):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    dash = p.add_run("— ")
    dash.font.name = "Arial"
    dash.font.size = Pt(7.5)
    dash.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    if bold_label:
        bl = p.add_run(bold_label + " ")
        bl.font.name = "Arial"
        bl.font.size = Pt(7.5)
        bl.font.bold = True
    t = p.add_run(str(text))
    t.font.name = "Arial"
    t.font.size = Pt(7.5)

def kv(cell, key, value):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    k = p.add_run(key + "  ")
    k.font.name = "Arial"
    k.font.size = Pt(7.5)
    k.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    v = p.add_run(str(value))
    v.font.name = "Arial"
    v.font.size = Pt(7.5)
    if value == "Not stated":
        v.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
        v.font.italic = True

def spacer(cell, pts=5):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(" ")
    r.font.size = Pt(pts)

def divider(cell):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    add_para_border_bottom(p, "E0E0E0", "2")

def metric_mini_table(cell, metrics, col_width_inches):
    n = len(metrics)
    col_w_dxa = int((col_width_inches * 1440) / n)
    tbl = cell.add_table(rows=1, cols=n)
    tbl.style = "Table Grid"
    tbl.width = int(col_width_inches * 1440)
    for i, (label, value) in enumerate(metrics):
        c = tbl.rows[0].cells[i]
        add_cell_shading(c, "F4F4F4")
        c.width = col_w_dxa
        lp = c.add_paragraph()
        lp.paragraph_format.space_before = Pt(3)
        lp.paragraph_format.space_after = Pt(1)
        lr = lp.add_run(label.upper())
        lr.font.name = "Arial"
        lr.font.size = Pt(6.5)
        lr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        vp = c.add_paragraph()
        vp.paragraph_format.space_before = Pt(0)
        vp.paragraph_format.space_after = Pt(3)
        vr = vp.add_run(str(value))
        vr.font.name = "Arial"
        vr.font.size = Pt(8.5)
        vr.font.bold = True
        if value == "Not stated":
            vr.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
            vr.font.italic = True

# ── WORD DOC BUILDER ─────────────────────────────────────────────────────────

def build_doc(d):
    doc = Document()

    sec = doc.sections[0]
    sec.page_width    = Inches(8.5)
    sec.page_height   = Inches(11)
    sec.left_margin   = sec.right_margin = Inches(0.5)
    sec.top_margin    = Inches(0.45)
    sec.bottom_margin = Inches(0.4)

    CW    = 7.5
    HALF  = CW / 2
    THIRD = CW / 3

    # ── HEADER ──
    ht = doc.add_table(rows=2, cols=1)
    ht.style = "Table Grid"
    ht.width = int(CW * 1440)

    tr = ht.rows[0].cells[0]
    add_cell_shading(tr, "EFEFEF")
    tr.width = int(CW * 1440)

    p = tr.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(ns(d.get("deal_name"), "Deal Name Not Stated"))
    r.font.name = "Arial"; r.font.size = Pt(13); r.font.bold = True

    p2 = tr.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(1)
    parts = [d.get("address"), d.get("city_state"), d.get("submarket")]
    ac = d.get("asset_class")
    if ac: parts.append(f"Class {ac}")
    sub_text = "  ·  ".join([str(x) for x in parts if x])
    r2 = p2.add_run(sub_text)
    r2.font.name = "Arial"; r2.font.size = Pt(8.5)
    r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    p3 = tr.add_paragraph()
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after = Pt(3)
    badges = [d.get("deal_type"), d.get("deal_status"), d.get("broker"), "All figures per OM · Not underwritten"]
    badge_text = "  |  ".join([str(x) for x in badges if x])
    r3 = p3.add_run(badge_text)
    r3.font.name = "Arial"; r3.font.size = Pt(7)
    r3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    r3.font.italic = True

    sr = ht.rows[1].cells[0]
    add_cell_shading(sr, "EFEFEF")

    stats = [
        ("Units",           ns(d.get("units"), "—")),
        ("Avg SF",          ns(d.get("avg_sf"), "—")),
        ("Yr Built / Reno", f"{ns(d.get('year_built'),'—')} / {ns(d.get('year_renovated'),'—')}"),
        ("Occupancy",       ns(d.get("physical_occupancy"), "Not stated")),
        ("Purchase Price",  ns(d.get("purchase_price"), "Not stated")),
        ("Price / Unit",    ns(d.get("price_per_unit"), "Not stated")),
        ("Going-In Cap",    ns(d.get("going_in_cap_rate"), "Not stated")),
    ]
    stat_tbl = sr.add_table(rows=1, cols=7)
    stat_tbl.style = "Table Grid"
    stat_tbl.width = int(CW * 1440)
    col_w = int((CW * 1440) / 7)
    for i, (label, value) in enumerate(stats):
        sc = stat_tbl.rows[0].cells[i]
        add_cell_shading(sc, "EFEFEF")
        sc.width = col_w
        lp = sc.add_paragraph()
        lp.paragraph_format.space_before = Pt(4)
        lp.paragraph_format.space_after = Pt(1)
        lr = lp.add_run(label.upper())
        lr.font.name = "Arial"; lr.font.size = Pt(6.5)
        lr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        vp = sc.add_paragraph()
        vp.paragraph_format.space_before = Pt(0)
        vp.paragraph_format.space_after = Pt(4)
        vr = vp.add_run(str(value))
        vr.font.name = "Arial"; vr.font.size = Pt(8.5); vr.font.bold = True
        if value in ("Not stated", "—"):
            vr.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
            vr.font.italic = True

    gap = doc.add_paragraph()
    gap.paragraph_format.space_before = Pt(0)
    gap.paragraph_format.space_after = Pt(3)

    # ── BODY: 2 columns ──
    bt = doc.add_table(rows=1, cols=2)
    bt.style = "Table Grid"
    bt.width = int(CW * 1440)
    L = bt.rows[0].cells[0]
    R = bt.rows[0].cells[1]
    for cell in (L, R):
        remove_cell_borders(cell)
        cell.width = int(HALF * 1440)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    section_header(L, "Investment Thesis")
    for b in (d.get("investment_thesis") or ["Not stated in OM"]):
        if b: bullet(L, b)
    spacer(L)

    section_header(L, "Business Plan")
    for b in (d.get("business_plan") or ["Not stated in OM"]):
        if b: bullet(L, b)
    spacer(L)

    section_header(L, "Property Summary")
    kv(L, "Construction:", ns(d.get("construction_type"), "Not stated"))
    kv(L, "Parking:", ns(d.get("parking"), "Not stated"))
    kv(L, "Stories:", ns(d.get("stories"), "Not stated"))
    kv(L, "Econ Occ:", ns(d.get("economic_occupancy"), "Not stated"))
    if d.get("amenities"):
        kv(L, "Amenities:", d["amenities"])
    if d.get("unit_mix"):
        kv(L, "Unit mix:", d["unit_mix"])
    spacer(L)

    section_header(L, "Location & Demand Drivers")
    for b in (d.get("location_bullets") or ["Not stated in OM"]):
        if b: bullet(L, b)

    section_header(R, "Pricing & Capex")
    metric_mini_table(R, [
        ("Purchase Price", ns(d.get("purchase_price"), "Not stated")),
        ("Price / Unit",   ns(d.get("price_per_unit"), "Not stated")),
    ], HALF - 0.12)
    spacer(R, 3)
    metric_mini_table(R, [
        ("Capex Total",  ns(d.get("capex_total"), "Not stated")),
        ("Capex / Unit", ns(d.get("capex_per_unit"), "Not stated")),
    ], HALF - 0.12)
    for b in (d.get("capex_bullets") or []):
        if b: bullet(R, b)
    spacer(R)

    section_header(R, "In-Place vs Pro Forma")
    metric_mini_table(R, [
        ("In-Place Rent",  ns(d.get("in_place_rent"), "Not stated")),
        ("Pro Forma Rent", ns(d.get("pro_forma_rent"), "Not stated")),
        ("Loss-to-Lease",  ns(d.get("loss_to_lease"), "Not stated")),
    ], HALF - 0.12)
    spacer(R, 3)

    t12 = d.get("t12_basis") or "T-12"
    kv(R, f"{t12} EGI:",  ns(d.get("t12_egi"), "Not stated"))
    kv(R, f"{t12} OpEx:", f"{ns(d.get('t12_opex'),'Not stated')}  ({ns(d.get('t12_opex_pct'),'—')})")
    kv(R, f"{t12} NOI:",  f"{ns(d.get('t12_noi'),'Not stated')}  ({ns(d.get('t12_noi_margin'),'—')} margin)")
    divider(R)

    stab = d.get("stab_label") or "Pro Forma"
    kv(R, f"{stab} EGI:",  ns(d.get("stab_egi"), "Not stated"))
    kv(R, f"{stab} OpEx:", f"{ns(d.get('stab_opex'),'Not stated')}  ({ns(d.get('stab_opex_pct'),'—')})")
    kv(R, f"{stab} NOI:",  f"{ns(d.get('stab_noi'),'Not stated')}  ({ns(d.get('stab_noi_margin'),'—')} margin)")
    spacer(R)

    section_header(R, "Capital Structure (As Stated in OM)")
    for label, key in [
        ("Lender / Program:", "lender"),
        ("Type:",             "debt_type"),
        ("Term / IO:",        "term_io"),
        ("Rate:",             "rate"),
        ("LTC / LTV:",        "ltc_ltv"),
        ("Equity:",           "equity"),
    ]:
        kv(R, label, ns(d.get(key), "Not stated"))
    divider(R)

    metric_mini_table(R, [
        ("Levered IRR",     ns(d.get("levered_irr"), "Not stated")),
        ("Equity Multiple", ns(d.get("equity_multiple"), "Not stated")),
        ("Avg CoC",         ns(d.get("avg_coc"), "Not stated")),
    ], HALF - 0.12)
    spacer(R, 3)
    kv(R, "Exit yr:",  ns(d.get("exit_year"), "Not stated"))
    kv(R, "Exit cap:", ns(d.get("exit_cap"), "Not stated"))
    if d.get("tax_notes"):
        divider(R)
        kv(R, "Tax notes:", d["tax_notes"])

    gap2 = doc.add_paragraph()
    gap2.paragraph_format.space_before = Pt(0)
    gap2.paragraph_format.space_after = Pt(3)

    bot = doc.add_table(rows=1, cols=3)
    bot.style = "Table Grid"
    bot.width = int(CW * 1440)
    bot_w = int(THIRD * 1440)

    risks_c   = bot.rows[0].cells[0]
    mitig_c   = bot.rows[0].cells[1]
    process_c = bot.rows[0].cells[2]

    for cell in (risks_c, mitig_c, process_c):
        add_cell_shading(cell, "F4F4F4")
        cell.width = bot_w
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    section_header(risks_c, "Key Risks")
    for b in (d.get("key_risks") or ["Not stated in OM"]):
        if b: bullet(risks_c, b)

    section_header(mitig_c, "Why This Works")
    for b in (d.get("why_this_works") or ["Not stated in OM"]):
        if b: bullet(mitig_c, b)

    section_header(process_c, "Process & Status")
    for label, key in [
        ("Broker:",          "broker"),
        ("Guidance:",        "guidance"),
        ("Bid / IC date:",   "bid_date"),
        ("Tours:",           "tour_status"),
        ("Internal status:", "internal_status"),
        ("Notes:",           "notes"),
    ]:
        kv(process_c, label, ns(d.get(key), "Not stated"))

    return doc

# ── PIPELINE ─────────────────────────────────────────────────────────────────

def extract_text(file_bytes):
    text = ""
    with pdfplumber.open(file_bytes) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                text += t + "\n"
    return text

def call_claude(pdf_text):
    truncated = pdf_text[:90000]
    client = anthropic.Anthropic(api_key=st.secrets["API_KEY"])
    msg = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=4000,
        messages=[{"role": "user", "content": EXTRACTION_PROMPT + truncated}]
    )
    raw = msg.content[0].text.strip()
    raw = re.sub(r"^```json\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw)
    return json.loads(raw)

# ── UI ────────────────────────────────────────────────────────────────────────

uploaded_file = st.file_uploader("Upload Offering Memorandum (PDF)", type="pdf")

if uploaded_file:
    if st.button("Generate 1-Pager", type="primary", use_container_width=True):

        with st.spinner("Reading PDF..."):
            pdf_bytes = io.BytesIO(uploaded_file.read())
            pdf_text = extract_text(pdf_bytes)

        with st.spinner("Extracting deal data..."):
            try:
                data = call_claude(pdf_text)
            except json.JSONDecodeError as e:
                st.error(f"Could not parse response as JSON: {e}")
                st.stop()
            except Exception as e:
                st.error(f"Error: {e}")
                st.stop()

        with st.spinner("Building Word document..."):
            doc = build_doc(data)
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)

        deal_name = data.get("deal_name") or "deal"
        filename = re.sub(r"[^\w\s-]", "", deal_name).strip().replace(" ", "_") + "_1pager.docx"

        st.success("Done.")
        st.download_button(
            label="Download Word Doc",
            data=buf,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

        with st.expander("View extracted JSON"):
            st.json(data)
else:
    st.info("Upload an OM PDF to get started.")
